"""
Step 1: Extract text from various document formats
Modular extraction for PDF, DOCX, and XLSX files.
Stores extracted text in data/processed/ ready for LLM processing.
"""

import pdfplumber
import json
from pathlib import Path
from datetime import datetime
from typing import Optional
import openpyxl
from docx import Document


# =====================================================
# PDF EXTRACTION
# =====================================================
def extract_pdf(pdf_path: Path) -> dict:
    """
    Extract text from a PDF file using pdfplumber.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Dictionary containing extracted text and metadata
    """
    extracted_data = {
        "source_file": str(pdf_path.name),
        "file_type": "pdf",
        "extraction_timestamp": datetime.now().isoformat(),
        "pages": [],
        "full_text": ""
    }
    
    all_text = []
    
    with pdfplumber.open(pdf_path) as pdf:
        extracted_data["total_pages"] = len(pdf.pages)
        
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            
            extracted_data["pages"].append({
                "page_number": page_num,
                "text": page_text,
                "char_count": len(page_text)
            })
            
            all_text.append(page_text)
    
    extracted_data["full_text"] = "\n\n".join(all_text)
    extracted_data["total_char_count"] = len(extracted_data["full_text"])
    
    print(f"   📄 PDF: {pdf_path.name}")
    print(f"      - Pages: {extracted_data['total_pages']}")
    print(f"      - Characters: {extracted_data['total_char_count']}")
    
    return extracted_data


# =====================================================
# DOCX EXTRACTION
# =====================================================
def extract_docx(docx_path: Path) -> dict:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary containing extracted text and metadata
    """
    extracted_data = {
        "source_file": str(docx_path.name),
        "file_type": "docx",
        "extraction_timestamp": datetime.now().isoformat(),
        "paragraphs": [],
        "tables": [],
        "full_text": ""
    }
    
    doc = Document(docx_path)
    all_text = []
    
    # Extract paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            extracted_data["paragraphs"].append({
                "index": i,
                "text": text,
                "style": para.style.name if para.style else None
            })
            all_text.append(text)
    
    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "table_index": table_idx,
            "rows": []
        }
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data["rows"].append(row_data)
            all_text.append(" | ".join(row_data))
        
        extracted_data["tables"].append(table_data)
    
    extracted_data["full_text"] = "\n".join(all_text)
    extracted_data["total_paragraphs"] = len(extracted_data["paragraphs"])
    extracted_data["total_tables"] = len(extracted_data["tables"])
    extracted_data["total_char_count"] = len(extracted_data["full_text"])
    
    print(f"   📝 DOCX: {docx_path.name}")
    print(f"      - Paragraphs: {extracted_data['total_paragraphs']}")
    print(f"      - Tables: {extracted_data['total_tables']}")
    print(f"      - Characters: {extracted_data['total_char_count']}")
    
    return extracted_data


# =====================================================
# XLSX EXTRACTION
# =====================================================
def extract_xlsx(xlsx_path: Path) -> dict:
    """
    Extract data from an XLSX file using openpyxl.
    
    Args:
        xlsx_path: Path to the XLSX file
        
    Returns:
        Dictionary containing extracted data and metadata
    """
    extracted_data = {
        "source_file": str(xlsx_path.name),
        "file_type": "xlsx",
        "extraction_timestamp": datetime.now().isoformat(),
        "sheets": [],
        "full_text": ""
    }
    
    all_text = []
    
    workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
    
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        sheet_data = {
            "sheet_name": sheet_name,
            "rows": [],
            "headers": []
        }
        
        rows_list = list(sheet.iter_rows(values_only=True))
        
        if rows_list:
            # First row as headers
            headers = [str(cell) if cell is not None else "" for cell in rows_list[0]]
            sheet_data["headers"] = headers
            
            # Process all rows
            for row_idx, row in enumerate(rows_list):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                sheet_data["rows"].append({
                    "row_index": row_idx,
                    "values": row_values
                })
                
                # Create text representation
                row_text = " | ".join([v for v in row_values if v])
                if row_text:
                    all_text.append(f"[{sheet_name}] {row_text}")
        
        extracted_data["sheets"].append(sheet_data)
    
    extracted_data["full_text"] = "\n".join(all_text)
    extracted_data["total_sheets"] = len(extracted_data["sheets"])
    extracted_data["total_char_count"] = len(extracted_data["full_text"])
    
    print(f"   📊 XLSX: {xlsx_path.name}")
    print(f"      - Sheets: {extracted_data['total_sheets']}")
    print(f"      - Characters: {extracted_data['total_char_count']}")
    
    return extracted_data


# =====================================================
# UNIFIED EXTRACTION
# =====================================================
def extract_file(file_path: Path) -> Optional[dict]:
    """
    Extract content from a file based on its extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dictionary with extracted content, or None if unsupported
    """
    ext = file_path.suffix.lower()
    
    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        return extract_xlsx(file_path)
    else:
        print(f"   ⚠️ Unsupported file type: {ext}")
        return None


def save_extraction(extracted_data: dict, output_dir: Path) -> tuple[Path, Path]:
    """
    Save extracted data as JSON and plain text.
    
    Args:
        extracted_data: Dictionary containing extracted content
        output_dir: Directory to save outputs
        
    Returns:
        Tuple of (json_path, txt_path)
    """
    source_name = Path(extracted_data["source_file"]).stem
    
    # Save JSON
    json_path = output_dir / f"{source_name}_extracted.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(extracted_data, f, indent=2, ensure_ascii=False)
    
    # Save plain text
    txt_path = output_dir / f"{source_name}_extracted.txt"
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(extracted_data["full_text"])
    
    return json_path, txt_path


# =====================================================
# MAIN PIPELINE
# =====================================================
def main():
    project_root = Path(__file__).parent.parent
    raw_dir = project_root / "data" / "raw"
    output_dir = project_root / "data" / "processed"
    
    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print("=" * 60)
    print("DOCUMENT EXTRACTION PIPELINE")
    print("=" * 60)
    
    # Files to extract
    files_to_extract = [
        raw_dir / "regulation.pdf",
        raw_dir / "policy.docx",
        raw_dir / "system_rules.xlsx"
    ]
    
    results = []
    
    for file_path in files_to_extract:
        if not file_path.exists():
            print(f"\n⚠️ File not found: {file_path}")
            continue
        
        print(f"\n📂 Extracting: {file_path.name}")
        
        # Extract content
        extracted = extract_file(file_path)
        
        if extracted:
            # Save outputs
            json_path, txt_path = save_extraction(extracted, output_dir)
            print(f"      ✓ Saved: {json_path.name}")
            print(f"      ✓ Saved: {txt_path.name}")
            
            results.append({
                "source": file_path.name,
                "json_output": str(json_path),
                "txt_output": str(txt_path),
                "char_count": extracted["total_char_count"]
            })
    
    # Summary
    print("\n" + "=" * 60)
    print("EXTRACTION COMPLETE")
    print("=" * 60)
    print(f"\n📊 Summary:")
    print(f"   Files processed: {len(results)}")
    for r in results:
        print(f"   - {r['source']}: {r['char_count']} chars")
    
    print(f"\n✅ All outputs saved to: {output_dir}")
    
    return results


if __name__ == "__main__":
    main()
